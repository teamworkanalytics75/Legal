"""Group chat workflow for legal document analysis and grading."""

from __future__ import annotations

import copy
import io
import logging
import os
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, Iterable, List, Callable

try:
    from autogen import AssistantAgent, GroupChat, GroupChatManager, UserProxyAgent
except ImportError:
    # Fallback for different AutoGen versions
    from autogen_agentchat.agents import AssistantAgent, UserProxyAgent
    from autogen_agentchat.teams import GroupChat, GroupChatManager

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).resolve().parent / "analysis_outputs"


def _is_termination_message(message: Dict[str, Any]) -> bool:
    """Return True if the message signals conversation termination."""
    content = message.get("content") if isinstance(message, dict) else None
    if isinstance(content, str):
        return content.strip().endswith("TERMINATE")
    return False


def _default_llm_config() -> Dict[str, Any]:
    """Build a default LLM configuration using environment variables."""
    model = os.environ.get("AUTOGEN_AGENT_MODEL", "gpt-4o-mini")
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        # Try loading from config file
        try:
            repo_root = Path(__file__).resolve().parents[2]
            config_path = repo_root / "config" / "OpenaiConfig.json"
            if config_path.exists():
                import json
                with open(config_path, 'r') as f:
                    config = json.load(f)
                    api_key = config.get("openai_api_key")
        except Exception:
            pass
    if not api_key:
        # Try .openai_api_key.txt as last resort
        repo_root = Path(__file__).resolve().parents[1]
        key_file = repo_root / ".openai_api_key.txt"
        if key_file.exists():
            candidate = key_file.read_text(encoding="utf-8").strip()
            if candidate:
                api_key = candidate
    if not api_key:
        raise RuntimeError(
            "OpenAI API key not configured. Set OPENAI_API_KEY, populate config/OpenaiConfig.json, or create .openai_api_key.txt."
        )
    
    return {
        "config_list": [
            {
                "model": model,
                "api_key": api_key,
                "api_type": "openai",
            }
        ]
    }


def _build_agents() -> tuple[
    AssistantAgent,
    AssistantAgent,
    AssistantAgent,
    AssistantAgent,
    Dict[str, Any],
]:
    """Create the planner and assistant agents for the group chat."""
    llm_config = _default_llm_config()
    planner = AssistantAgent(
        name="Planner",
        system_message=(
            "You coordinate a collaborative legal review.\n"
            "First turn: summarize the key objectives for examining the provided document, "
            "highlight specific sections or schedules to focus on, and hand off explicitly to the Grader.\n"
            "Later turns (only if needed): reconcile feedback and request targeted follow-ups.\n"
            "Always reference prior agent contributions and avoid producing final scores yourself."
        ),
        llm_config=copy.deepcopy(llm_config),
        is_termination_msg=_is_termination_message,
    )
    grader = AssistantAgent(
        name="Grader",
        system_message=(
            "You evaluate legal quality.\n"
            "Wait for the Planner to set priorities. When you respond, reference the Planner's guidance "
            "and the source document. Provide numeric scores with rationale and highlight the strongest "
            "and weakest elements. Conclude by inviting the Critic to respond next."
        ),
        llm_config=copy.deepcopy(llm_config),
        is_termination_msg=_is_termination_message,
    )
    critic = AssistantAgent(
        name="Critic",
        system_message=(
            "You refine the work after the Grader.\n"
            "Incorporate the Planner's objectives and the Grader's findings. Identify omissions, factual "
            "gaps, or clarity issues and propose targeted revisions. Reference previous messages and "
            "pass the baton to the Probabilist when done."
        ),
        llm_config=copy.deepcopy(llm_config),
        is_termination_msg=_is_termination_message,
    )
    probabilist = AssistantAgent(
        name="Probabilist",
        system_message=(
            "You conclude the review with likelihood estimates.\n"
            "Wait until the Critic has finished. Integrate the prior discussion to estimate the probability "
            "that each major factual claim is true, estimate the likelihood that the matter qualifies as a U.S. national "
            "security concern (state a percentage and a short justification referencing evidence), and summarize residual "
            "risks. End your response with the word 'TERMINATE' to close the session."
        ),
        llm_config=copy.deepcopy(llm_config),
        is_termination_msg=_is_termination_message,
    )
    return planner, grader, critic, probabilist, copy.deepcopy(llm_config)


def _load_local_document(path: Path) -> str:
    """Load text from a local document."""
    if path.is_dir():
        raise RuntimeError(f"{path} is a directory. Provide a path to a document file.")
    suffix = path.suffix.lower()
    if suffix == ".txt":
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="utf-8-sig")
    if suffix == ".docx":
        try:
            from docx import Document # type: ignore
        except ImportError as exc: # pragma: no cover - dependency shim
            raise RuntimeError("python-docx must be installed to read .docx files.") from exc
        document = Document(path)
        return "\n".join(par.text for par in document.paragraphs)
    if suffix == ".pdf":
        try:
            import fitz # type: ignore
        except ImportError as exc: # pragma: no cover - dependency shim
            raise RuntimeError("PyMuPDF must be installed to read .pdf files.") from exc
        with fitz.open(path) as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig")
    except PermissionError as exc:
        raise RuntimeError(f"Permission denied when reading {path}.") from exc


def _download_drive_file(file_id: str) -> tuple[bytes, str]:
    """Download a Google Drive file and return its bytes and MIME type."""
    try:
        from googleapiclient.discovery import build # type: ignore
        from googleapiclient.http import MediaIoBaseDownload # type: ignore
        from google.oauth2 import service_account # type: ignore
    except ImportError as exc: # pragma: no cover - optional dependency
        raise RuntimeError("google-api-python-client is required for Drive access.") from exc

    credentials_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
    if not credentials_path:
        raise RuntimeError("Set GOOGLE_APPLICATION_CREDENTIALS to a service account JSON file for Drive access.")
    service_account_path = Path(credentials_path)
    if not service_account_path.exists():
        raise RuntimeError(f"Service account file not found: {service_account_path}")

    scopes = ["https://www.googleapis.com/auth/drive.readonly"]
    credentials = service_account.Credentials.from_service_account_file(str(service_account_path), scopes=scopes)
    service = build("drive", "v3", credentials=credentials, cache_discovery=False)

    metadata = service.files().get(fileId=file_id, fields="id, name, mimeType").execute()
    mime_type = metadata.get("mimeType", "")

    def _download_request() -> Any:
        if mime_type == "application/vnd.google-apps.document":
            return service.files().export_media(
                fileId=file_id,
                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        return service.files().get_media(fileId=file_id)

    request = _download_request()
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buffer.seek(0)

    if mime_type == "application/vnd.google-apps.document":
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return buffer.read(), mime_type


def _parse_drive_payload(data: bytes, mime_type: str) -> str:
    """Convert a downloaded Drive payload into text."""
    if mime_type in {"text/plain"}:
        return data.decode("utf-8")
    if mime_type in {"application/pdf"}:
        try:
            import fitz # type: ignore
        except ImportError as exc: # pragma: no cover - dependency shim
            raise RuntimeError("PyMuPDF must be installed to read Drive PDF content.") from exc
        with fitz.open(stream=data, filetype="pdf") as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    if mime_type in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}:
        try:
            from docx import Document # type: ignore
        except ImportError as exc: # pragma: no cover - dependency shim
            raise RuntimeError("python-docx must be installed to read Drive DOCX content.") from exc
        document = Document(io.BytesIO(data))
        return "\n".join(par.text for par in document.paragraphs)
    raise RuntimeError(f"Unsupported Google Drive MIME type: {mime_type}")


def load_document(path_or_id: str) -> str:
    """Load a document from a local path or Google Drive file ID."""
    candidate_path = Path(path_or_id)
    if candidate_path.exists():
        return _load_local_document(candidate_path)
    data, mime_type = _download_drive_file(path_or_id)
    return _parse_drive_payload(data, mime_type)


def _pysmile_available() -> bool:
    """Check whether PySMILE is available."""
    try:
        import importlib.util

        return importlib.util.find_spec("pysmile") is not None
    except Exception: # pragma: no cover - defensive
        return False


def _load_bn_context() -> str:
    """Load BN context when requested and available."""
    use_bn = os.environ.get("USE_BN", "true").lower() == "true"
    if not use_bn:
        logger.info("USE_BN disabled. Running in analysis-only mode.")
        return ""
    if not _pysmile_available():
        logger.info("PySMILE not available. Skipping BN context.")
        return ""
    module_name = "wizardweb_stable"
    module_path = Path(__file__).resolve().parent.parent / "experiments" / "WizardWeb1.1.4_STABLE.py"
    if not module_path.exists():
        logger.warning("WizardWeb BN module not found. Skipping BN context.")
        return ""
    try:
        import importlib.util

        spec = importlib.util.spec_from_file_location(module_name, module_path)
        if spec is None or spec.loader is None:
            logger.error("Could not load WizardWeb module specification.")
            return ""
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
    except Exception as exc:
        logger.error("Failed to import WizardWeb module: %s", exc)
        return ""

    bn_query = getattr(module, "bn_query_direct", None)
    if not callable(bn_query):
        logger.error("bn_query_direct not available in WizardWeb module.")
        return ""
    logger.info("Including BN context from WizardWeb model.")
    return str(bn_query({}))


def _build_prompt(document_text: str, bn_context: str | None) -> str:
    """Create the joint analysis prompt."""
    header: List[str] = [
        "Collaborative Evaluation Protocol:",
        "1. Planner acts first: outline a concise evaluation plan referencing document sections, then hand off to the Grader.",
        "2. Grader responds once: score the document using the Planner's plan, citing evidence, and invite the Critic to refine.",
        "3. Critic addresses the Grader's findings: propose revisions or risk mitigations, then request the Probabilist's assessment.",
        "4. Probabilist concludes: integrate earlier feedback, estimate success probabilities, evaluate whether the matter is likely to be treated as a U.S. national security concern (with justification), summarize residual risks, and end with 'TERMINATE'.",
        "Each agent must acknowledge previous contributions, avoid repeating full rubrics, and keep the conversation focused.",
        "",
    ]
    if bn_context:
        header.extend(["[BN Context]", bn_context.strip(), ""])
    header.extend(["[Document]", document_text.strip()])
    return "\n".join(header)


def _normalize_content(content: Any) -> str:
    """Convert message content to plain text."""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        fragments: List[str] = []
        for item in content:
            if isinstance(item, dict) and "text" in item:
                fragments.append(str(item["text"]))
            elif isinstance(item, str):
                fragments.append(item)
        return "\n".join(fragments).strip()
    return str(content).strip()


def _collect_agent_messages(groupchat: GroupChat, agent_names: Iterable[str]) -> List[Dict[str, str]]:
    """Collect messages from the specified agents."""
    agent_set = set(agent_names)
    collected: List[Dict[str, str]] = []
    for message in getattr(groupchat, "messages", []):
        sender = message.get("name") or message.get("role") or ""
        if sender not in agent_set:
            continue
        text = _normalize_content(message.get("content", ""))
        if not text:
            continue
        collected.append({"sender": sender, "content": text})
    return collected


def _select_bottom_line(messages: List[Dict[str, str]]) -> str:
    """Select the bottom-line message for the report."""
    for record in reversed(messages):
        if record["sender"] == "Grader":
            return record["content"]
    return messages[-1]["content"] if messages else "No agent analysis captured."


def _write_report(
    document_path: str,
    use_bn: bool,
    messages: List[Dict[str, str]],
    output_dir: Path,
) -> Path:
    """Write a DOCX report summarizing the chat analysis."""
    from docx import Document # type: ignore

    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = Path(document_path).stem or "analysis"
    report_path = output_dir / f"{base_name}_analysis_{timestamp}.docx"

    doc = Document()
    doc.add_heading("Group Analysis Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Source document: {document_path}")
    doc.add_paragraph(f"Bayesian context enabled: {'Yes' if use_bn else 'No'}")

    doc.add_heading("Bottom Line Up Front", level=2)
    doc.add_paragraph(_select_bottom_line(messages))

    doc.add_heading("Agent Contributions", level=2)
    by_agent: Dict[str, List[str]] = {}
    for message in messages:
        by_agent.setdefault(message["sender"], []).append(message["content"])

    for agent_name, entries in by_agent.items():
        doc.add_heading(agent_name, level=3)
        for entry in entries:
            doc.add_paragraph(entry, style="List Bullet")

    doc.save(report_path)
    return report_path


def main() -> None:
    """Entry point for the group analysis chat."""
    try:
        document_identifier = input("Enter path or Drive ID of document:\n> ").strip()
        if not document_identifier:
            raise RuntimeError("A document path or Drive ID is required.")

        print(f"[Info] Loading document: {document_identifier}")
        document_text = load_document(document_identifier)
        print(f"[Info] Document loaded successfully ({len(document_text)} characters)")

        print("[Info] Loading BN context...")
        bn_context = _load_bn_context()
        if bn_context:
            print("[Info] BN results included in context.")
        else:
            print("[Info] Running without BN context.")

        print("[Info] Building agents...")
        planner, grader, critic, probabilist, llm_config = _build_agents()
        print("[Info] Agents created successfully")
        
        print("[Info] Starting group chat...")
        groupchat = GroupChat(
            agents=[planner, grader, critic, probabilist],
            messages=[],
            max_round=12,
            speaker_selection_method="round_robin",
            allow_repeat_speaker=False,
            send_introductions=True,
        )
        manager = GroupChatManager(
            groupchat,
            llm_config=llm_config,
            is_termination_msg=_is_termination_message,
        )

        prompt = _build_prompt(document_text, bn_context if bn_context else "")
        messages = [{"role": "user", "content": prompt}]
        
        print("[Info] Running analysis...")
        success, _ = manager.run_chat(messages=messages, sender=probabilist, config=groupchat)
        if not success:
            print("[Warn] Group chat terminated without a successful completion.")
        else:
            print("[Info] Analysis completed successfully!")

        agent_transcript = _collect_agent_messages(
            groupchat,
            [planner.name, grader.name, critic.name, probabilist.name],
        )
        if not agent_transcript:
            print("[Warn] No agent messages captured for reporting.")
        else:
            report_path = _write_report(
                document_identifier,
                bool(bn_context),
                agent_transcript,
                OUTPUT_DIR,
            )
            print(f"[Info] Report saved to {report_path}")
            
    except Exception as e:
        print(f"[Error] {e}")
        print("\nTroubleshooting tips:")
        print(" - Make sure the document path is correct")
        print(" - Check that you have the required packages installed")
        print(" - Verify your OpenAI API key is set")
        raise


if __name__ == "__main__":
    main()
