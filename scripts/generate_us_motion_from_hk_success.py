#!/usr/bin/env python3
"""
Generate US Motion to Seal from HK Success Case

Uses:
1. Successful HK Motion to Seal (as template/precedent)
2. HK Statement of Claim (as facts/evidence)
3. Full pipeline: CatBoost Analysis → Semantic Kernel → Autogen → Validation
4. Target: CatBoost confidence > 75% for US Motion

This demonstrates using successful foreign cases to guide US motion drafting.
"""

import asyncio
import json
import logging
import sys
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

# Add paths
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root / "writer_agents" / "code"))

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not available. Install with: pip install python-docx")

from WorkflowOrchestrator import Conductor
from workflow_config import WorkflowStrategyConfig
# Backward compatibility alias
WorkflowStrategyExecutor = Conductor
from sk_config import SKConfig
from insights import CaseInsights, EvidenceItem
from agents import ModelConfig

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def extract_docx_text(docx_path: Path) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Install with: pip install python-docx")

    if not docx_path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    doc = Document(docx_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    full_text = "\n\n".join(text_parts)
    logger.info(f"Extracted {len(full_text)} characters from {docx_path.name}")
    return full_text


async def generate_us_motion_from_hk_success(
    hk_motion_path: Path,
    hk_statement_path: Path,
    target_confidence: float = 0.75,
    autogen_model: str = "qwen2.5:14b",
    sk_use_local: bool = True,
    max_iterations: int = 3
) -> Dict[str, Any]:
    """
    Generate US Motion to Seal using HK success case as precedent.

    Args:
        hk_motion_path: Path to successful HK Motion to Seal DOCX
        hk_statement_path: Path to HK Statement of Claim TXT
        target_confidence: Target CatBoost confidence (default 0.75 = 75%)
        autogen_model: Model for Autogen
        sk_use_local: Use local LLMs
        max_iterations: Max workflow iterations
    """

    logger.info("="*70)
    logger.info("GENERATE US MOTION FROM HK SUCCESS CASE")
    logger.info("="*70)

    # Load documents
    logger.info(f"\nLoading HK Motion to Seal: {hk_motion_path}")
    hk_motion_text = extract_docx_text(hk_motion_path)

    logger.info(f"\nLoading HK Statement of Claim: {hk_statement_path}")
    with open(hk_statement_path, 'r', encoding='utf-8') as f:
        hk_statement_text = f.read()

    logger.info(f"\nDocuments loaded:")
    logger.info(f"  HK Motion: {len(hk_motion_text)} characters")
    logger.info(f"  HK Statement: {len(hk_statement_text)} characters")

    # Create case insights combining both documents
    case_summary = f"""
    Generate a US Motion to Seal based on:

    1. SUCCESSFUL HK PRECEDENT: A successful Hong Kong Motion to Seal (provided below)
       - This motion was granted in Hong Kong
       - Use its structure, arguments, and persuasive elements as guidance
       - Adapt for US legal context (28 U.S.C. Section 1782, federal courts)

    2. FACTUAL BASIS: Statement of Claim (provided below)
       - Contains allegations that constitute a US national security matter
       - Plaintiff is a US citizen
       - Involves Harvard-China connections, PRC leadership references, endangerment

    OBJECTIVE: Draft a US Motion to Seal that:
    - Adapts successful HK motion structure to US legal framework
    - Incorporates Statement of Claim facts
    - Meets CatBoost success criteria (target confidence: {target_confidence:.0%})
    - Addresses US national security implications

    HK MOTION TO SEAL (SUCCESSFUL PRECEDENT):
    {hk_motion_text[:3000]}...

    STATEMENT OF CLAIM (FACTUAL BASIS):
    {hk_statement_text[:2000]}...
    """

    insights = CaseInsights(
        reference_id=f"US-MOTION-FROM-HK-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        summary=case_summary,
        jurisdiction="US Federal Court, D. Massachusetts",
        evidence=[
            EvidenceItem(
                node_id="hk_successful_motion",
                state="provided",
                description=f"Successful HK Motion to Seal ({len(hk_motion_text)} chars) - use as precedent"
            ),
            EvidenceItem(
                node_id="hk_statement_of_claim",
                state="provided",
                description=f"HK Statement of Claim ({len(hk_statement_text)} chars) - factual basis"
            ),
            EvidenceItem(
                node_id="national_security_matter",
                state="confirmed",
                description="Allegations constitute US national security matter"
            ),
            EvidenceItem(
                node_id="target_confidence",
                state=str(target_confidence),
                description=f"Target CatBoost confidence: {target_confidence:.0%}"
            )
        ],
        posteriors=[],
        case_style="Motion for Leave to File Under Seal Under 28 U.S.C. Section 1782"
    )

    # Configure workflow
    if sk_use_local:
        autogen_config = ModelConfig(use_local=True, local_model="qwen2.5:14b")
    else:
        autogen_config = ModelConfig(model=autogen_model)

    sk_config = SKConfig(use_local=sk_use_local)

    workflow_config = WorkflowStrategyConfig(
        autogen_config=autogen_config,
        sk_config=sk_config,
        max_iterations=max_iterations,
        enable_sk_planner=True,
        enable_quality_gates=True,
        auto_commit_threshold=target_confidence,  # Use target confidence as threshold
        enable_autogen_review=True,
        validation_strict=True,
        enable_chroma_integration=True,
        chroma_persist_directory="./chroma_collections"
    )

    # Initialize conductor
    executor = Conductor(config=workflow_config)

    # Run workflow
    logger.info("\nStarting hybrid workflow (Autogen → SK → CatBoost)...")
    logger.info(f"Target confidence: {target_confidence:.0%}")
    logger.info("Using LOCAL LLMs" if sk_use_local else "Using OpenAI models")

    try:
        deliverable = await executor.run_hybrid_workflow(insights)

        # Extract final document
        final_document = deliverable.edited_document or (
            deliverable.sections[0].body if deliverable.sections else ""
        )

        # Get validation results
        validation_results = {}
        if hasattr(executor, '_last_state') and executor._last_state:
            validation_results = executor._last_state.validation_results or {}

        # Check CatBoost confidence
        catboost_confidence = validation_results.get("catboost_validation", {}).get("predicted_success_probability", 0.0)
        meets_target = catboost_confidence >= target_confidence

        # Compile results
        results = {
            "success": True,
            "final_motion": final_document,
            "target_confidence": target_confidence,
            "actual_confidence": catboost_confidence,
            "meets_target": meets_target,
            "validation_results": validation_results,
            "hk_motion_preview": hk_motion_text[:500],
            "hk_statement_preview": hk_statement_text[:500],
            "workflow_metadata": {
                "iterations": executor.iteration_controller.get_iteration_summary() if hasattr(executor, 'iteration_controller') else {},
                "deliverable_sections": len(deliverable.sections) if deliverable.sections else 0
            }
        }

        logger.info("\n" + "="*70)
        logger.info("MOTION GENERATION COMPLETE")
        logger.info("="*70)
        logger.info(f"Target Confidence: {target_confidence:.1%}")
        logger.info(f"Actual Confidence: {catboost_confidence:.1%}")
        logger.info(f"Meets Target: {'YES' if meets_target else 'NO'}")
        logger.info(f"Document Length: {len(final_document)} characters")

        return results

    except Exception as e:
        logger.error(f"Motion generation failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "final_motion": "",
            "validation_results": {}
        }


def main():
    """Command-line interface."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate US Motion to Seal from successful HK Motion"
    )

    parser.add_argument(
        "--hk-motion",
        type=str,
        required=True,
        help="Path to HK Motion to Seal DOCX file"
    )
    parser.add_argument(
        "--hk-statement",
        type=str,
        default=None,
        help="Path to HK Statement of Claim TXT file (default: auto-detect)"
    )
    parser.add_argument(
        "--target-confidence",
        type=float,
        default=0.75,
        help="Target CatBoost confidence (default: 0.75 = 75%%)"
    )
    parser.add_argument(
        "--max-iterations",
        type=int,
        default=3,
        help="Maximum workflow iterations (default: 3)"
    )
    parser.add_argument(
        "--output",
        type=str,
        help="Output JSON file path"
    )
    parser.add_argument(
        "--motion-output",
        type=str,
        help="Output motion text file path"
    )

    args = parser.parse_args()

    # Resolve paths
    hk_motion_path = Path(args.hk_motion)
    if not hk_motion_path.exists():
        print(f"ERROR: HK Motion file not found: {hk_motion_path}")
        return

    # Auto-detect statement if not provided
    if args.hk_statement:
        hk_statement_path = Path(args.hk_statement)
    else:
        # Try default location
        default_statement = project_root / "case_law_data" / "tmp_corpus" / "Exhibit 2 — Certified Statement of Claim (Hong Kong, 2 Jun 2025).txt"
        if default_statement.exists():
            hk_statement_path = default_statement
            print(f"Using auto-detected statement: {hk_statement_path}")
        else:
            print(f"ERROR: Statement file not found. Please provide --hk-statement")
            return

    if not hk_statement_path.exists():
        print(f"ERROR: Statement file not found: {hk_statement_path}")
        return

    # Run generation
    results = asyncio.run(generate_us_motion_from_hk_success(
        hk_motion_path=hk_motion_path,
        hk_statement_path=hk_statement_path,
        target_confidence=args.target_confidence,
        sk_use_local=True,  # Use local LLMs
        max_iterations=args.max_iterations
    ))

    # Save results
    output_dir = project_root / "outputs"
    output_dir.mkdir(exist_ok=True)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = output_dir / f"us_motion_from_hk_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"\nResults saved to: {output_path}")

    # Save motion text
    if args.motion_output:
        motion_path = Path(args.motion_output)
    else:
        motion_path = output_path.with_suffix('.motion.txt')

    with open(motion_path, 'w', encoding='utf-8') as f:
        f.write(results.get("final_motion", ""))
    print(f"Motion text saved to: {motion_path}")

    # Print summary
    if results.get("success"):
        print("\n" + "="*70)
        print("SUCCESS - US MOTION GENERATED")
        print("="*70)
        print(f"Target Confidence: {results['target_confidence']:.1%}")
        print(f"Actual Confidence: {results['actual_confidence']:.1%}")
        print(f"Meets Target: {'YES' if results['meets_target'] else 'NO'}")
        print(f"\nMotion Preview:")
        print("-"*70)
        print(results.get("final_motion", "")[:1000] + "..." if len(results.get("final_motion", "")) > 1000 else results.get("final_motion", ""))
    else:
        print("\n" + "="*70)
        print("ERROR - GENERATION FAILED")
        print("="*70)
        print(f"Error: {results.get('error', 'Unknown error')}")


if __name__ == "__main__":
    main()
